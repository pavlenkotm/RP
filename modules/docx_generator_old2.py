"""
Модуль для генерации Word документов расчётов на прочность
"""
import os
import shutil
import re
from typing import Dict, Tuple, Optional
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import win32com.client
from PIL import Image


class DOCXGenerator:
    """Класс для генерации документов Word"""
    
    def __init__(self, template_path: str, output_dir: str, config_manager, logger):
        """
        Инициализация генератора документов
        
        Args:
            template_path: путь к шаблону Word
            output_dir: директория для сохранения документов
            config_manager: менеджер конфигурации
            logger: логгер
        """
        self.template_path = template_path
        self.output_dir = output_dir
        self.config = config_manager
        self.logger = logger
        
        # Создаём выходную директорию, если её нет
        if not os.path.exists(output_dir):
            os.makedirs(output_dir, exist_ok=True)
        
        # Оригинальные значения из шаблона (для замены)
        self.template_values = {
            'title_name': 'Змейка без песочницы',
            'title_article': '810152',
            'text_article': 'GA8808',
            'text_category': 'песочниц',  # для склонения
            'children_count': '10',
            'mass_child': '32.5',
        }
    
    def generate_document(self, product_data: Dict, technical_data: Dict[str, Tuple[str, str]]) -> str:
        """
        Генерация документа для изделия
        
        Args:
            product_data: данные изделия из Excel
            technical_data: технические данные из паспорта
            
        Returns:
            Путь к созданному документу
        """
        # Формируем имя выходного файла
        article = product_data['article'].replace('/', '-').replace('\\', '-')
        name = product_data['name'][:50]  # Ограничиваем длину
        name = self._sanitize_filename(name)
        
        output_filename = f"{article}_{name}_РП.docx"
        output_path = os.path.join(self.output_dir, output_filename)
        
        # Копируем шаблон
        shutil.copy(self.template_path, output_path)
        
        # Открываем документ
        doc = Document(output_path)
        
        # Заполняем документ
        self._process_document(doc, product_data, technical_data)
        
        # Сохраняем документ
        doc.save(output_path)
        
        # Обновляем поля через Word COM (если доступен)
        try:
            self._update_fields_com(output_path)
        except Exception as e:
            self.logger.log_warning(f"Не удалось обновить поля через COM: {str(e)}")
        
        return output_path
    
    def _process_document(self, doc: Document, product_data: Dict, technical_data: Dict[str, Tuple[str, str]]):
        """
        Обработка документа - замена текста, удаление картинок, вставка нового изображения
        """
        # 1. Подготовка данных для замены
        replacements = self._prepare_replacements(product_data, technical_data)
        
        # 2. Замена текста во всех параграфах
        for paragraph in doc.paragraphs:
            self._replace_in_paragraph(paragraph, replacements)
        
        # 3. Замена текста в таблицах
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._replace_in_paragraph(paragraph, replacements)
        
        # 4. Удаление всех изображений кроме "Рис. X - общий вид"
        self._remove_images_except_main(doc)
        
        # 5. Вставка нового изображения
        if product_data.get('image_path') and os.path.exists(product_data['image_path']):
            self._insert_main_image(doc, product_data['image_path'])
    
    def _prepare_replacements(self, product_data: Dict, technical_data: Dict) -> Dict[str, str]:
        """
        Подготовка словаря замен
        
        Args:
            product_data: данные изделия
            technical_data: технические данные
            
        Returns:
            Словарь {старый_текст: новый_текст}
        """
        article = product_data['article']
        name = product_data['name']
        category = product_data['category']
        children_count = product_data['children_count']
        
        # Параметры нагрузок
        mass_child = self.config.get_mass_child()  # 53.8 кг
        total_mass = children_count * mass_child
        
        # Расчёт сил от детей
        # Fh = 0.2 * M * g (горизонтальная)
        # Fz = M * g (вертикальная)
        g = 10  # м/с² (упрощённо)
        Fh = 0.2 * total_mass * g
        Fz = total_mass * g
        
        # Определяем склонение категории для текста
        category_gen = self._get_category_genitive(category)  # родительный падеж
        
        # Словарь замен
        replacements = {
            # Титульный лист
            'Змейка без песочницы арт.810152': f'{name} арт.{article}',
            'Змейка без песочницы': name,
            'арт.810152': f'арт.{article}',
            '810152': article,
            
            # В тексте - артикулы
            'GA8808': article,
            f'артикул GA8808': f'артикул {article}',
            
            # В тексте - названия и категории
            'песочницы, артикул GA8808': f'{category_gen}, артикул {article}',
            'песочницы': category_gen,
            
            # Нагрузки от детей
            '10 детей': f'{children_count} детей',
            f'{children_count} детей' if children_count > 1 else f'{children_count} ребенка': 
                f'{children_count} детей' if children_count > 4 or children_count == 0 else 
                f'{children_count} ребенка' if children_count == 1 else f'{children_count} детей',
            
            '32.5 кг': f'{mass_child:.1f} кг',
            '32,5 кг': f'{mass_child:.1f} кг',
            
            # Силы (пересчитываем)
            'Fh = 646,8 Н': f'Fh = {Fh:.1f} Н',
            'Fh = 646.8 Н': f'Fh = {Fh:.1f} Н',
            'Fz = 6468 Н': f'Fz = {Fz:.0f} Н',
            'Fz = 6468.0 Н': f'Fz = {Fz:.0f} Н',
            
            # В заключении
            'песочницы, артикул GA8808': f'{category_gen}, артикул {article}',
        }
        
        return replacements
    
    @staticmethod
    def _get_category_genitive(category: str) -> str:
        """
        Получить категорию в родительном падеже
        
        Args:
            category: категория в именительном падеже
            
        Returns:
            Категория в родительном падеже
        """
        genitive_map = {
            'Домики': 'игрового домика',
            'Игровые комплексы': 'игрового комплекса',
            'Игровые элементы': 'игрового элемента',
            'Мини-беседки': 'мини-беседки',
            'Беседки': 'беседки',
            'Песочницы': 'песочницы'
        }
        return genitive_map.get(category, 'конструкции')
    
    @staticmethod
    def _param_to_placeholder(param_name: str) -> str:
        """
        Преобразование названия параметра в плейсхолдер
        
        Args:
            param_name: название параметра
            
        Returns:
            Плейсхолдер вида {PARAM_NAME}
        """
        # Упрощённое преобразование - берём первые слова
        words = param_name.upper().split()[:3]
        key = '_'.join(words)
        # Убираем спецсимволы
        key = ''.join(c if c.isalnum() or c == '_' else '' for c in key)
        return f"{{{key}}}"
    
    @staticmethod
    def _replace_in_paragraph(paragraph, replacements: Dict[str, str]):
        """
        Замена текста в параграфе с сохранением форматирования
        
        Args:
            paragraph: параграф документа
            replacements: словарь замен
        """
        # Получаем полный текст параграфа
        full_text = paragraph.text
        
        # Проверяем, нужна ли замена
        modified = False
        for old_text, new_text in replacements.items():
            if old_text in full_text:
                full_text = full_text.replace(old_text, new_text)
                modified = True
        
        # Если были изменения, обновляем текст
        if modified:
            # Сохраняем форматирование первого run
            if paragraph.runs:
                # Очищаем все runs
                for run in paragraph.runs[1:]:
                    run.text = ''
                # Обновляем первый run
                paragraph.runs[0].text = full_text
            else:
                paragraph.text = full_text
    
    def _remove_images_except_main(self, doc: Document):
        """
        Удаление всех изображений кроме места для основного изображения
        
        Args:
            doc: объект документа
        """
        # Проходим по всем параграфам и удаляем изображения
        # Оставляем только параграф с текстом "Рис" для вставки нового изображения
        for paragraph in doc.paragraphs:
            # Проверяем, есть ли в параграфе изображение
            for run in paragraph.runs:
                # Ищем элементы изображений
                drawings = run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing')
                for drawing in drawings:
                    # Проверяем, это не параграф с "Рис"
                    if 'Рис' not in paragraph.text:
                        # Удаляем изображение
                        parent = drawing.getparent()
                        parent.remove(drawing)
    
    def _insert_main_image(self, doc: Document, image_path: str):
        """
        Вставка основного изображения изделия
        
        Args:
            doc: объект документа
            image_path: путь к изображению
        """
        try:
            # Ищем параграф для вставки изображения
            # Обычно это параграф с текстом "Рис. X" или рядом с "Общий вид"
            target_paragraph = None
            
            for i, paragraph in enumerate(doc.paragraphs):
                text = paragraph.text.lower()
                # Ищем параграф с "рис" и "общ" (общий вид)
                if ('рис' in text or 'общ' in text) and ('вид' in text or '.' in text):
                    # Проверяем, нет ли уже изображения в этом параграфе
                    has_image = False
                    for run in paragraph.runs:
                        drawings = run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing')
                        if drawings:
                            has_image = True
                            break
                    
                    if not has_image:
                        target_paragraph = paragraph
                        break
            
            # Если не нашли подходящий параграф, ищем первое упоминание "Рис"
            if not target_paragraph:
                for paragraph in doc.paragraphs:
                    if 'Рис' in paragraph.text or 'рис' in paragraph.text:
                        target_paragraph = paragraph
                        break
            
            # Если всё ещё не нашли, создаём новый параграф после раздела "ОБЩИЕ СВЕДЕНИЯ"
            if not target_paragraph:
                for i, paragraph in enumerate(doc.paragraphs):
                    if 'ОБЩИЕ СВЕДЕНИЯ' in paragraph.text:
                        if i + 3 < len(doc.paragraphs):
                            target_paragraph = doc.paragraphs[i + 3]
                        break
            
            # Вставляем изображение
            if target_paragraph:
                # Открываем изображение для получения размеров
                img = Image.open(image_path)
                width, height = img.size
                
                # Максимальная ширина (в дюймах)
                max_width = 6.0
                
                # Масштабируем с сохранением пропорций
                if width > height:
                    new_width = max_width
                    new_height = (height / width) * max_width
                else:
                    new_height = max_width
                    new_width = (width / height) * max_width
                
                # Очищаем параграф от текста (если он пустой или содержит только "Рис.")
                if len(target_paragraph.text.strip()) < 20:
                    target_paragraph.clear()
                
                # Добавляем изображение
                run = target_paragraph.add_run()
                run.add_picture(image_path, width=Inches(new_width))
                target_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
        except Exception as e:
            self.logger.log_warning(f"Ошибка вставки изображения: {str(e)}")
    
    @staticmethod
    def _update_fields_com(doc_path: str):
        """
        Обновление полей документа через COM (PAGE, NUMPAGES, TOC)
        
        Args:
            doc_path: путь к документу
        """
        try:
            # Создаём COM объект Word
            word = win32com.client.DispatchEx("Word.Application")
            word.Visible = False
            
            # Открываем документ
            doc = word.Documents.Open(os.path.abspath(doc_path))
            
            # Обновляем все поля
            doc.Fields.Update()
            
            # Обновляем содержание (TOC)
            for toc in doc.TablesOfContents:
                toc.Update()
            
            # Сохраняем и закрываем
            doc.Save()
            doc.Close()
            word.Quit()
        
        except Exception as e:
            raise Exception(f"Ошибка обновления полей через COM: {str(e)}")
    
    @staticmethod
    def _sanitize_filename(filename: str) -> str:
        """
        Очистка имени файла от недопустимых символов
        
        Args:
            filename: исходное имя файла
            
        Returns:
            Очищенное имя файла
        """
        # Недопустимые символы для Windows
        invalid_chars = '<>:"/\\|?*'
        for char in invalid_chars:
            filename = filename.replace(char, '_')
        return filename.strip()
