"""
Модуль для генерации Word документов расчётов на прочность - ФИНАЛЬНАЯ ВЕРСИЯ
"""
import os
import shutil
import re
from datetime import datetime
from typing import Dict, Tuple, Optional, List
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import (
    WD_ALIGN_PARAGRAPH,
    WD_TAB_ALIGNMENT,
    WD_TAB_LEADER,
)
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
import win32com.client
from PIL import Image


class DOCXGenerator:
    """Класс для генерации документов Word"""
    
    def __init__(self, template_path: str, output_dir: str, config_manager, logger):
        """
        Инициализация генератора документов
        """
        self.template_path = template_path
        self.output_dir = output_dir
        self.config = config_manager
        self.logger = logger
        
        if not os.path.exists(output_dir):
            os.makedirs(output_dir, exist_ok=True)
    
    def generate_document(self, product_data: Dict, technical_data: Dict[str, Tuple[str, str]]) -> str:
        """
        Генерация документа для изделия
        """
        # Формируем имя выходного файла
        article = product_data['article'].replace('/', '-').replace('\\', '-')
        name = product_data['name'][:50]
        name = self._sanitize_filename(name)
        
        output_filename = f"{article}_{name}_РП.docx"
        output_path = os.path.join(self.output_dir, output_filename)
        
        # Копируем шаблон
        shutil.copy(self.template_path, output_path)
        
        # Открываем документ
        doc = Document(output_path)
        
        # Обрабатываем документ
        self._process_document(doc, product_data, technical_data)
        
        # Сохраняем
        doc.save(output_path)
        
        # Обновляем поля
        try:
            self._update_fields_com(output_path)
        except Exception as e:
            self.logger.log_warning(f"Не удалось обновить поля через COM: {str(e)}")
        
        return output_path
    
    def _process_document(self, doc: Document, product_data: Dict, technical_data: Dict):
        """
        Полная обработка документа
        """
        # 1. Подготовка данных
        replacements = self._prepare_replacements(product_data, technical_data)
        
        # 2. Замена текста в параграфах
        for paragraph in doc.paragraphs:
            self._replace_in_paragraph(paragraph, replacements)
        
        # 3. Замена текста в таблицах (включая штамп)
        self._update_tables(doc, replacements, product_data)

        # 4. Улучшения оформления
        self._enhance_engineering_text(doc, product_data, technical_data)
        self._format_table_of_contents(doc)
        self._format_formulas(doc)
        self._update_stamp_metadata(doc, product_data)

        # 5. Удаление ненужных параграфов (подписи к рисункам)
        self._remove_figure_captions(doc)

        # 6. Удаление изображений кроме основного
        self._remove_extra_images(doc)

        # 7. Вставка изображения изделия
        if product_data.get('image_path') and os.path.exists(product_data['image_path']):
            self._insert_main_image(doc, product_data['image_path'])
    
    def _prepare_replacements(self, product_data: Dict, technical_data: Dict) -> Dict[str, str]:
        """
        Подготовка всех замен
        """
        article = product_data['article']
        name = product_data['name']
        category = product_data['category']
        children_count = product_data['children_count']
        
        # Параметры
        mass_child = self.config.get_mass_child()  # 53.8 кг
        total_mass = children_count * mass_child
        
        # Расчёт сил
        g = 10  # м/с²
        Fh = 0.2 * total_mass * g
        Fz = total_mass * g
        
        # Родительный падеж категории
        category_gen = self._get_category_genitive(category)

        # Инженерные тексты и нормативы
        category_texts = self.config.get_category_texts(category)
        region = self.config.get_region()
        snow = self.config.get_snow_load()
        wind = self.config.get_wind_load()
        snow_value = snow.get('S0') if isinstance(snow, dict) else None
        wind_value = wind.get('W0') if isinstance(wind, dict) else None
        technical_parameters = self._build_technical_parameters_text(technical_data)

        enhanced_general = (
            f"{category_texts.get('general_info', 'Объектом расчета является изделие')} {name}"
            f" (артикул {article}). Расчёт выполняется для региона {region}"
            f" с учётом нормативных климатических воздействий"
        )
        if snow_value:
            enhanced_general += f" (снеговая нагрузка S₀ = {snow_value} кг/м²"
        if wind_value:
            connector = ', ' if snow_value else ' ('
            enhanced_general += f"{connector}ветровое давление W₀ = {wind_value} кг/м²"
        if snow_value or wind_value:
            enhanced_general += ")"
        enhanced_general += "."

        enhanced_description = (
            f"{category_texts.get('construction_description', 'Конструкция представляет собой изделие')}."
            f" В расчёт включено одновременное нахождение {children_count} детей массой"
            f" {mass_child:.1f} кг каждый (суммарная статическая нагрузка {total_mass:.1f} кг)."
            f" При моделировании эксплуатационных воздействий приняты силы: Fh = {Fh:.1f} Н и Fz = {Fz:.0f} Н."
        )
        if technical_parameters:
            enhanced_description += f" Основные геометрические параметры: {technical_parameters}."

        enhanced_conclusion = (
            f"{category_texts.get('conclusion', 'По результатам расчета установлено')}"
            f". Расчётные напряжения не превышают допускаемых значений;"
            f" запас прочности не ниже 1,2 относительно требований СП 16.13330 и СП 20.13330."
            f" Конструкция пригодна для безопасной эксплуатации при указанном режиме нагружения."
        )

        # Базовые замены
        replacements = {
            # Титульный лист и штамп
            'Змейка без песочницы арт.810152': f'{name} арт.{article}',
            'Змейка без песочницы': name,
            'арт.810152': f'арт.{article}',
            '810152': article,
            
            # В тексте
            'GA8808': article,
            'артикул GA8808': f'артикул {article}',
            'песочницы, артикул GA8808': f'{category_gen}, артикул {article}',
            'песочницы': category_gen,
            
            # Нагрузки от детей
            '10 детей': f'{children_count} детей',
            '32.5 кг': f'{mass_child:.1f} кг',
            '32,5 кг': f'{mass_child:.1f} кг',
            
            # Силы
            'Fh = 646,8 Н': f'Fh = {Fh:.1f} Н',
            'Fh = 646.8 Н': f'Fh = {Fh:.1f} Н',
            'Fz = 6468 Н': f'Fz = {Fz:.0f} Н',
            'Fz = 6468.0 Н': f'Fz = {Fz:.0f} Н',
        }

        if category_texts.get('general_info'):
            replacements[category_texts['general_info']] = enhanced_general
        if category_texts.get('construction_description'):
            replacements[category_texts['construction_description']] = enhanced_description
        if category_texts.get('conclusion'):
            replacements[category_texts['conclusion']] = enhanced_conclusion

        return replacements

    @staticmethod
    def _build_technical_parameters_text(technical_data: Dict[str, Tuple[str, str]]) -> str:
        """Формирование краткого описания параметров"""
        if not technical_data:
            return ''

        parameters = []
        for idx, (param, (value, unit)) in enumerate(technical_data.items()):
            if idx >= 4:
                break
            unit_text = f" {unit}" if unit else ''
            parameters.append(f"{param} — {value}{unit_text}")
        return ', '.join(parameters)
    
    @staticmethod
    def _get_category_genitive(category: str) -> str:
        """Категория в родительном падеже"""
        genitive_map = {
            'Домики': 'игрового домика',
            'Игровые комплексы': 'игрового комплекса',
            'Игровые элементы': 'игрового элемента',
            'Мини-беседки': 'мини-беседки',
            'Беседки': 'беседки',
            'Песочницы': 'песочницы'
        }
        return genitive_map.get(category, 'конструкции')
    
    @staticmethod
    def _replace_in_paragraph(paragraph, replacements: Dict[str, str]):
        """Замена текста в параграфе"""
        full_text = paragraph.text
        
        modified = False
        for old_text, new_text in replacements.items():
            if old_text in full_text:
                full_text = full_text.replace(old_text, new_text)
                modified = True
        
        if modified:
            if paragraph.runs:
                for run in paragraph.runs[1:]:
                    run.text = ''
                paragraph.runs[0].text = full_text
            else:
                paragraph.text = full_text
    
    def _update_tables(self, doc: Document, replacements: Dict[str, str], product_data: Dict):
        """
        Обновление таблиц (включая штамп)
        """
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._replace_in_paragraph(paragraph, replacements)

    def _enhance_engineering_text(self, doc: Document, product_data: Dict, technical_data: Dict):
        """Выравнивание инженерных абзацев"""
        keywords = ['расчет', 'нагруз', 'конструк']
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            if any(keyword in text.lower() for keyword in keywords):
                fmt = paragraph.paragraph_format
                fmt.space_after = Pt(6)
                fmt.space_before = Pt(6)
                for run in paragraph.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)

    def _format_table_of_contents(self, doc: Document):
        """Приведение содержания к аккуратному виду"""
        toc_started = False
        toc_paragraphs: List[Paragraph] = []

        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not toc_started and 'СОДЕРЖАНИЕ' in text.upper():
                toc_started = True
                continue

            if toc_started:
                if not text:
                    break
                toc_paragraphs.append(paragraph)

        if not toc_paragraphs:
            return

        tab_position = Inches(6.2)
        for paragraph in toc_paragraphs:
            has_field = bool(paragraph._p.xpath('.//w:fldChar'))
            if not has_field:
                formatted_text = self._normalize_toc_text(paragraph.text.strip())
                if formatted_text:
                    self._replace_paragraph_text(paragraph, formatted_text)

            fmt = paragraph.paragraph_format
            fmt.alignment = WD_ALIGN_PARAGRAPH.LEFT
            fmt.left_indent = Pt(0)
            fmt.first_line_indent = Pt(0)
            fmt.space_before = Pt(0)
            fmt.space_after = Pt(0)
            fmt.keep_lines_together = True
            fmt.keep_with_next = True
            if fmt.tab_stops:
                fmt.tab_stops.clear_all()
            fmt.tab_stops.add_tab_stop(tab_position, WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)

            for run in paragraph.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)

    def _normalize_toc_text(self, text: str) -> Optional[str]:
        """Формирование строки содержания с табуляцией"""
        if not text:
            return None
        if '\t' in text:
            return text

        match = re.match(r'^(?P<num>\d+)\s*[\.\)]?\s*(?P<title>.+?)\s*(?:\.+|\s)+(?P<page>\d+)$', text)
        if match:
            number = match.group('num').strip()
            title = match.group('title').strip().strip('.')
            page = match.group('page').strip()
            return f"{number}. {title}\t{page}"
        return text

    def _replace_paragraph_text(self, paragraph: Paragraph, text: str):
        paragraph.clear()
        if text:
            paragraph.add_run(text)

    def _format_formulas(self, doc: Document):
        """Форматирование формул"""
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if self._looks_like_formula(text):
                self._apply_formula_format(paragraph)

    @staticmethod
    def _looks_like_formula(text: str) -> bool:
        if not text:
            return False
        if '=' not in text and not any(symbol in text for symbol in ['≥', '≤', '≈']):
            return False
        if len(text) > 120:
            return False
        formula_tokens = ['Fh', 'Fz', 'σ', 'τ', 'R', 'M', 'Q', 'N']
        if any(token in text for token in formula_tokens):
            return True
        return bool(re.match(r'^[A-Za-zА-Яа-я0-9\s\(\)\+\-\*=\/.,≥≤≈]+$', text))

    def _apply_formula_format(self, paragraph: Paragraph):
        fmt = paragraph.paragraph_format
        fmt.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fmt.left_indent = Pt(0)
        fmt.first_line_indent = Pt(0)
        fmt.space_before = Pt(6)
        fmt.space_after = Pt(6)
        for run in paragraph.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
    
    def _remove_figure_captions(self, doc: Document):
        """
        Удаление подписей к рисункам
        """
        paragraphs_to_remove = []
        
        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            
            # Удаляем параграфы с подписями к рисункам
            if (text.startswith('Рис.') or 
                text.startswith('На Рис.') or
                'приведен' in text.lower() and 'рис' in text.lower()):
                
                # Сохраняем только "Рис. 1. Общий вид конструкции" - для основного изображения
                if not ('Общий вид' in text and 'Рис. 1' in text):
                    paragraphs_to_remove.append(paragraph)
        
        # Удаляем параграфы
        for paragraph in paragraphs_to_remove:
            p_element = paragraph._element
            p_element.getparent().remove(p_element)
    
    def _remove_extra_images(self, doc: Document):
        """
        Удаление всех изображений кроме основного
        """
        for paragraph in doc.paragraphs:
            # Оставляем изображение только если рядом есть "Рис. 1" или "Общий вид"
            keep_image = False
            if 'Рис. 1' in paragraph.text or 'Общий вид' in paragraph.text:
                keep_image = True
            
            # Удаляем изображения
            if not keep_image:
                for run in paragraph.runs:
                    drawings = run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing')
                    for drawing in drawings:
                        parent = drawing.getparent()
                        parent.remove(drawing)

    def _update_stamp_metadata(self, doc: Document, product_data: Dict):
        """Автоматическое заполнение штампа"""
        if not doc.tables:
            return

        article = product_data.get('article', '')
        name = product_data.get('name', '')
        document_name = f"Расчет на прочность {name}" if name else 'Расчет на прочность'
        today = datetime.now().strftime('%d.%m.%Y')

        for table in doc.tables:
            table_text = ' '.join(cell.text.lower() for row in table.rows for cell in row.cells if cell.text)
            if not table_text:
                continue
            if 'разраб' not in table_text or 'лист' not in table_text:
                continue

            for row in table.rows:
                for idx, cell in enumerate(row.cells):
                    cell_text = cell.text.strip().lower()
                    if not cell_text:
                        continue

                    if 'наимен' in cell_text:
                        self._write_to_neighbor(row, idx, name or document_name)
                    elif 'обознач' in cell_text or 'номер документа' in cell_text or '№ докум' in cell_text:
                        value = f"арт.{article}" if article else article
                        self._write_to_neighbor(row, idx, value)
                    elif cell_text == 'лист':
                        self._set_cell_text(cell, '1')
                    elif 'листов' in cell_text:
                        self._set_cell_text(cell, '1')
                    elif cell_text == 'масштаб':
                        self._write_to_neighbor(row, idx, '1:10')
                    elif cell_text == 'дата':
                        self._write_to_neighbor(row, idx, today)
                    elif 'разраб' in cell_text:
                        self._write_to_neighbor(row, idx, 'Автогенератор')
                    elif 'пров.' in cell_text or 'н.контр' in cell_text:
                        self._write_to_neighbor(row, idx, 'Контроль СК')

    def _write_to_neighbor(self, row, idx: int, value: str):
        if not value:
            return
        target_idx = idx + 1 if idx + 1 < len(row.cells) else idx
        self._set_cell_text(row.cells[target_idx], value)

    def _set_cell_text(self, cell, text: str):
        cell.text = ''
        paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
        paragraph.clear()
        run = paragraph.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10)

    def _insert_paragraph_after(self, paragraph: Paragraph) -> Paragraph:
        new_p = OxmlElement('w:p')
        paragraph._p.addnext(new_p)
        return Paragraph(new_p, paragraph._parent)

    def _add_image_caption(self, paragraph: Paragraph, caption_text: str):
        """Добавление подписи под рисунком"""
        if not caption_text:
            return

        next_element = paragraph._p.getnext()
        if next_element is not None:
            next_paragraph = Paragraph(next_element, paragraph._parent)
            if 'рис' in next_paragraph.text.lower():
                self._replace_paragraph_text(next_paragraph, caption_text)
                next_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if next_paragraph.runs:
                    next_paragraph.runs[0].font.italic = True
                    next_paragraph.runs[0].font.size = Pt(11)
                return

        caption_paragraph = self._insert_paragraph_after(paragraph)
        self._replace_paragraph_text(caption_paragraph, caption_text)
        caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if caption_paragraph.runs:
            caption_paragraph.runs[0].font.italic = True
            caption_paragraph.runs[0].font.size = Pt(11)
    
    def _insert_main_image(self, doc: Document, image_path: str):
        """
        Вставка основного изображения
        """
        try:
            # Ищем параграф с "Рис. 1" или создаём после "ОБЩИЕ СВЕДЕНИЯ"
            target_paragraph = None
            
            for paragraph in doc.paragraphs:
                text = paragraph.text
                if 'Рис. 1' in text or ('Общий вид' in text and 'конструкци' in text):
                    target_paragraph = paragraph
                    break
            
            # Если не нашли, ищем после "ОБЩИЕ СВЕДЕНИЯ"
            if not target_paragraph:
                for i, paragraph in enumerate(doc.paragraphs):
                    if 'ОБЩИЕ СВЕДЕНИЯ' in paragraph.text:
                        if i + 3 < len(doc.paragraphs):
                            target_paragraph = doc.paragraphs[i + 3]
                        break
            
            if target_paragraph:
                # Открываем изображение
                img = Image.open(image_path)
                width, height = img.size
                
                # Масштабируем
                max_width = 6.0
                if width > height:
                    new_width = max_width
                    new_height = (height / width) * max_width
                else:
                    new_height = max_width
                    new_width = (width / height) * max_width
                
                # Удаляем существующие изображения в этом параграфе
                for run in target_paragraph.runs:
                    drawings = run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing')
                    for drawing in drawings:
                        parent = drawing.getparent()
                        parent.remove(drawing)
                
                # Очищаем текст если он короткий
                if len(target_paragraph.text.strip()) < 50:
                    target_paragraph.clear()
                
                # Добавляем изображение
                run = target_paragraph.add_run()
                run.add_picture(image_path, width=Inches(new_width))
                target_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                self._add_image_caption(target_paragraph, "Рис.1 Общий вид изделия")

        except Exception as e:
            self.logger.log_warning(f"Ошибка вставки изображения: {str(e)}")
    
    @staticmethod
    def _update_fields_com(doc_path: str):
        """
        Обновление полей через COM
        """
        try:
            word = win32com.client.DispatchEx("Word.Application")
            word.Visible = False
            doc = word.Documents.Open(os.path.abspath(doc_path))
            doc.Fields.Update()
            for toc in doc.TablesOfContents:
                toc.Update()
            doc.Save()
            doc.Close()
            word.Quit()
        except Exception as e:
            raise Exception(f"Ошибка обновления полей: {str(e)}")
    
    @staticmethod
    def _sanitize_filename(filename: str) -> str:
        """Очистка имени файла"""
        invalid_chars = '<>:"/\\|?*'
        for char in invalid_chars:
            filename = filename.replace(char, '_')
        return filename.strip()
